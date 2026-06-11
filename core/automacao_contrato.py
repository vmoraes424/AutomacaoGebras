import sys
from pathlib import Path

# Permite executar: python core/automacao_contrato.py (além de python -m core.automacao_contrato)
_root = Path(__file__).resolve().parent.parent
if str(_root) not in sys.path:
    sys.path.insert(0, str(_root))

import time
import os
import traceback
import requests
import base64
import hashlib
import glob
from datetime import datetime, timedelta, timezone
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.run import Run
from docxtpl import DocxTemplate

from core.config import (
    CLICKSIGN_ACCESS_TOKEN,
    CLICKSIGN_BASE_URL,
    CLICKSIGN_RATE_LIMIT_BUFFER_SEC,
    CLICKSIGN_RATE_LIMIT_MAX_RETRIES,
    DEV_PULAR_CLICKSIGN,
    DEV_HUB_SEM_APROVACAO_PLUNE,
    AUTOMACAO_WORKER_ENABLED,
    FORMULARIO_WEB_ENABLED,
    INTERVALO_POLLING_SEGUNDOS,
    MODELO_DOCX,
    PASTA_SAIDA,
    PIPEDRIVE_API_TOKEN,
    PULAR_HUB,
    TESTE_PLUNE_SEM_ASSINATURA,
)
from core.form_deal_adapter import (
    atualizar_deal_form_status,
    listar_deal_ids_formulario_aguardando_worker,
    preparar_deal_para_automacao,
)
from core.database import (
    carregar_deals_processados,
    garantir_envelope_para_hub,
    salvar_deal_processado,
)
from core.envelope_state import (
    carregar_pedidos_plune_criados,
    buscar_por_deal_id,
    limpar_template_local_envelope,
    listar_aguardando_pedido_plune,
    marcar_pedidos_aprovados,
    salvar_envelope_pendente,
)
from core.plune_pedido import (
    PluneError,
    anexar_contrato_local_aos_pedidos_deal,
    aprovar_pedidos_plune,
    buscar_parceiro_plune_por_documento,
    criar_pedido_plune,
    extrair_numeros_pedidos_plune,
    formatar_linha_pedidos_plune_contrato,
    obter_numeros_pedidos_plune_deal,
)
from core.pipedrive_stages import marcar_deal_como_ganho, deal_elegivel_formulario_contrato
from core.pipedrive_validations import (
    DealValidationError,
    reabrir_deal_com_erros,
    reabrir_deal_falha_automacao,
)
from core.pipedrive_fields import (
    FIELD_CONTATO_CONTRATANTE,
    FIELD_CONTATO_FINANCEIRO,
    FIELD_DATA_IMPLANTACAO,
    FIELD_DATA_PRIMEIRA_COBRANCA,
    FIELD_DOCUMENTO,
    FIELD_ENDERECO,
    FIELD_CIDADE,
    FIELD_GESTAO_ACL,
    FIELD_GESTAO_USINA_FOTOVOLTAICA,
    FIELD_INDICADORES_QUALIDADE,
    FIELD_INSCRICAO_ESTADUAL,
    FIELD_INSCRICAO_MUNICIPAL,
    FIELD_NOTAS,
    FIELD_PERCENTUAL_EXITO,
    FIELD_QTD_SOLE,
    FIELD_QUALIDADE_ENERGIA,
    FIELD_VALOR_IMPLANTACAO,
    FIELD_VALOR_MENSAL,
    extrair_signatarios,
    signatarios_omitidos_por_email_duplicado,
    formatar_data_hora_brasilia,
    formatar_data_ptbr,
    formatar_quantidade_uc,
    get_enum_label,
    get_nome_cliente,
    get_numero_contrato,
    get_cidade_estado,
    get_contato_gestor_contrato,
    get_val,
    buscar_deal_por_id,
)
from core.pipedrive_files import baixar_docx_contrato_padrao_deal
from core.aviso_comercial import enviar_aviso_comercial_etapa1
from core.gebras_defaults import EMAIL_COMERCIAL_AUTOMACAO

# --- CONFIGURAÇÕES (ver config.py) ---
API_TOKEN = PIPEDRIVE_API_TOKEN

# Evita repetir o mesmo aviso sobre linha legada em deals_processados a cada ciclo
_deals_legado_ja_avisados: set[str] = set()


def _registrar_deal_processado_worker(deal_id: str, evento_str: str) -> None:
    salvar_deal_processado(deal_id, evento_str)
    try:
        atualizar_deal_form_status(deal_id, "processed")
    except Exception as exc:
        print(
            f"[!] Deal {deal_id}: falha ao marcar formulário como processed: {exc}",
            flush=True,
        )


# ---------- GERENCIAMENTO DE ESTADO ----------
def _parse_deal_timestamp_utc(timestamp_str: str) -> datetime | None:
    """ISO-8601 do Pipedrive (Z, offset, fração de segundos)."""
    if not timestamp_str or not isinstance(timestamp_str, str):
        return None
    s = timestamp_str.strip()
    if not s:
        return None
    if s.endswith("Z"):
        s = s[:-1] + "+00:00"
    try:
        dt = datetime.fromisoformat(s)
    except ValueError:
        return None
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    else:
        dt = dt.astimezone(timezone.utc)
    return dt


# ---------- FUNÇÕES AUXILIARES ----------
def formatar_moeda(valor):
    try:
        val = float(valor)
        return f"R$ {val:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except (ValueError, TypeError):
        return str(valor)


# ---------- CLICKSIGN CLIENT (V3) ----------


def _parse_clicksign_datetime(raw: str | None) -> datetime | None:
    if not raw:
        return None
    texto = str(raw).strip()
    if not texto:
        return None
    if texto.endswith("Z"):
        texto = texto[:-1] + "+00:00"
    try:
        return datetime.fromisoformat(texto)
    except ValueError:
        pass
    for fmt in ("%Y-%m-%dT%H:%M:%S", "%Y-%m-%d %H:%M:%S"):
        try:
            return datetime.strptime(texto[:19], fmt)
        except ValueError:
            continue
    return None


class ClicksignClient:
    def __init__(self, base_url: str, access_token: str):
        self.base_url = base_url.rstrip("/")
        self.access_token = access_token
        self.session = requests.Session()
        self.session.headers.update(
            {"Accept": "application/json", "Content-Type": "application/vnd.api+json"}
        )

    def _url(self, path: str):
        sep = "&" if "?" in path else "?"
        return f"{self.base_url}{path}{sep}access_token={self.access_token}"

    def _seconds_until_rate_reset(self, response: requests.Response) -> float | None:
        reset_raw = response.headers.get("X-Rate-Limit-Reset")
        if reset_raw:
            try:
                wait = int(reset_raw) - time.time() + CLICKSIGN_RATE_LIMIT_BUFFER_SEC
                return max(wait, 0.5)
            except ValueError:
                pass
        retry_after = response.headers.get("Retry-After")
        if retry_after:
            try:
                return max(float(retry_after) + CLICKSIGN_RATE_LIMIT_BUFFER_SEC, 0.5)
            except ValueError:
                pass
        return None

    def _request(self, method: str, path: str, ctx: str, **kwargs) -> requests.Response:
        url = self._url(path)
        response = None
        for attempt in range(1, CLICKSIGN_RATE_LIMIT_MAX_RETRIES + 1):
            response = self.session.request(method, url, **kwargs)
            if response.status_code != 429:
                return response

            wait = self._seconds_until_rate_reset(response)
            if wait is None:
                wait = min(2**attempt, 30)
            remaining = response.headers.get("X-Rate-Limit-Remaining", "?")
            reset_at = response.headers.get("X-Rate-Limit-Reset", "?")
            print(
                f"[!] Clicksign rate limit em {ctx} "
                f"(tentativa {attempt}/{CLICKSIGN_RATE_LIMIT_MAX_RETRIES}, "
                f"remaining={remaining}, reset={reset_at}). "
                f"Aguardando {wait:.1f}s..."
            )
            time.sleep(wait)

        assert response is not None
        self._raise(response, ctx)
        return response

    def _raise(self, r, ctx):
        if not r.ok:
            try:
                body = r.json()
            except Exception:
                body = r.text
            raise RuntimeError(f"[Clicksign] {ctx} -> {r.status_code}: {body}")

    def create_envelope(self, name: str):
        payload = {
            "data": {
                "type": "envelopes",
                "attributes": {
                    "name": name,
                    "locale": "pt-BR",
                    "auto_close": True,
                    "block_after_refusal": True,
                },
            }
        }
        r = self._request("POST", "/envelopes", "create_envelope", json=payload)
        self._raise(r, "create_envelope")
        return r.json()["data"]["id"]

    def upload_document_base64(self, envelope_id, file_path):
        print("[*] Upload documento (V3)...")
        with open(file_path, "rb") as f:
            raw = base64.b64encode(f.read()).decode()
        header = "data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,"
        payload = {
            "data": {
                "type": "documents",
                "attributes": {
                    "filename": os.path.basename(file_path),
                    "content_base64": header + raw,
                },
            }
        }
        r = self._request(
            "POST",
            f"/envelopes/{envelope_id}/documents",
            "upload_document",
            json=payload,
        )
        self._raise(r, "upload_document")
        return r.json()["data"]["id"]

    def create_signer(self, envelope_id, name, email, group_number):
        payload = {
            "data": {
                "type": "signers",
                "attributes": {
                    "name": name,
                    "email": email,
                    "group": group_number,
                    "communicate_events": {
                        "signature_request": "email",
                        "signature_reminder": "email",
                        "document_signed": "email",
                    },
                },
            }
        }
        r = self._request(
            "POST",
            f"/envelopes/{envelope_id}/signers",
            "create_signer",
            json=payload,
        )
        self._raise(r, "create_signer")
        return r.json()["data"]["id"]

    def create_sign_requirement(self, envelope_id, signer_id, document_id):
        payload = {
            "data": {
                "type": "requirements",
                "attributes": {"action": "agree", "role": "sign"},
                "relationships": {
                    "document": {"data": {"type": "documents", "id": document_id}},
                    "signer": {"data": {"type": "signers", "id": signer_id}},
                },
            }
        }
        r = self._request(
            "POST",
            f"/envelopes/{envelope_id}/requirements",
            "create_sign_requirement",
            json=payload,
        )
        self._raise(r, "create_sign_requirement")

    def create_auth_requirement(self, envelope_id, signer_id, document_id):
        payload = {
            "data": {
                "type": "requirements",
                "attributes": {"action": "provide_evidence", "auth": "email"},
                "relationships": {
                    "signer": {"data": {"type": "signers", "id": signer_id}},
                    "document": {"data": {"type": "documents", "id": document_id}},
                },
            }
        }
        r = self._request(
            "POST",
            f"/envelopes/{envelope_id}/requirements",
            "create_auth_requirement",
            json=payload,
        )
        self._raise(r, "create_auth_requirement")

    def activate_envelope(self, envelope_id):
        payload = {
            "data": {
                "id": envelope_id,
                "type": "envelopes",
                "attributes": {"status": "running"},
            }
        }
        r = self._request(
            "PATCH", f"/envelopes/{envelope_id}", "activate_envelope", json=payload
        )
        self._raise(r, "activate_envelope")

    def cancel_envelope(self, envelope_id: str) -> None:
        """
        Cancela um envelope em andamento.
        Regra da automação: ao regenerar (re-won antes da última assinatura),
        o envelope antigo deve ser invalidado para evitar aprovação por engano.
        """
        payload = {
            "data": {
                "id": envelope_id,
                "type": "envelopes",
                "attributes": {"status": "canceled"},
            }
        }
        r = self._request(
            "PATCH", f"/envelopes/{envelope_id}", "cancel_envelope", json=payload
        )
        # Não levanta em falha: pode já estar fechado/cancelado; seguimos com a regeneração.
        if not r.ok:
            try:
                body = r.json()
            except Exception:
                body = r.text
            print(
                f"[!] Clicksign: falha ao cancelar envelope {envelope_id} "
                f"(HTTP {r.status_code}): {body}",
                flush=True,
            )

    def notify_signer_manual(self, envelope_id, signer_id):
        payload = {
            "data": {
                "type": "notifications",
                "attributes": {"message": "Sua vez de assinar o contrato."},
            }
        }
        r = self._request(
            "POST",
            f"/envelopes/{envelope_id}/signers/{signer_id}/notifications",
            "notify_signer",
            json=payload,
        )
        self._raise(r, "notify_signer")

    def get_envelope_attributes(self, envelope_id: str) -> dict:
        r = self._request("GET", f"/envelopes/{envelope_id}", "get_envelope")
        self._raise(r, "get_envelope")
        return r.json().get("data", {}).get("attributes", {}) or {}

    def get_envelope_status(self, envelope_id: str) -> str | None:
        try:
            return self.get_envelope_attributes(envelope_id).get("status")
        except RuntimeError:
            return None

    def get_ultima_assinatura_data_ptbr(self, envelope_id: str) -> str:
        """
        Data da última assinatura (dd/mm/aaaa) para Venda.Pedido.DataEntrega.
        Usa o maior `modified` entre os signatários; fallback: `modified` do envelope fechado.
        """
        from core.pipedrive_fields import formatar_data_ptbr

        attrs = self.get_envelope_attributes(envelope_id)
        if attrs.get("status") != "closed":
            return formatar_data_ptbr(None)

        ultima: datetime | None = _parse_clicksign_datetime(attrs.get("modified"))
        r = self._request(
            "GET", f"/envelopes/{envelope_id}/signers", "list_signers"
        )
        if r.ok:
            signers = r.json().get("data") or []
            if isinstance(signers, dict):
                signers = [signers]
            for signer in signers:
                mod = _parse_clicksign_datetime(
                    (signer.get("attributes") or {}).get("modified")
                )
                if mod and (ultima is None or mod > ultima):
                    ultima = mod

        if ultima:
            return ultima.strftime("%d/%m/%Y")
        return formatar_data_ptbr(attrs.get("modified"))

    def baixar_pdf_assinado(self, envelope_id: str) -> tuple[bytes, str] | None:
        """Baixa o PDF assinado do primeiro documento do envelope (status closed)."""
        r = self._request(
            "GET", f"/envelopes/{envelope_id}/documents", "list_documents"
        )
        if not r.ok:
            return None
        documentos = r.json().get("data") or []
        if isinstance(documentos, dict):
            documentos = [documentos]
        for documento in documentos:
            doc_id = documento.get("id")
            if not doc_id:
                continue
            r_doc = self._request(
                "GET",
                f"/envelopes/{envelope_id}/documents/{doc_id}",
                "get_document",
            )
            if not r_doc.ok:
                continue
            data = r_doc.json().get("data", {}) or {}
            attrs = data.get("attributes", {}) or {}
            downloads = attrs.get("downloads") or {}
            url = downloads.get("signed_file_url") or downloads.get("signed_file")
            if not url:
                # Em alguns tenants Clicksign v3 o link vem em data.links.files.signed
                links = data.get("links", {}) or {}
                files = links.get("files", {}) or {}
                url = (
                    files.get("signed")
                    or files.get("signed_file_url")
                    or files.get("signed_file")
                )
            if not url:
                continue
            if str(url).startswith("/"):
                base = self.base_url.split("/api")[0]
                url = f"{base.rstrip('/')}{url}"
            nome = str(attrs.get("filename") or "contrato_assinado.pdf")
            if not nome.lower().endswith(".pdf"):
                nome = f"{Path(nome).stem}.pdf"
            resp = requests.get(url, timeout=180)
            if not resp.ok:
                continue
            return resp.content, nome
        return None


# ---------- FLUXO RÁPIDO (FIRE & FORGET COM GRUPOS) ----------
def clicksign_fire_and_forget(
    doc_path: str,
    envelope_name: str,
    sign_sequence: list,
    *,
    omitidos: list[dict] | None = None,
):
    cs = ClicksignClient(CLICKSIGN_BASE_URL, CLICKSIGN_ACCESS_TOKEN)

    print(f"[*] 1. Criando envelope '{envelope_name}'...")
    envelope_id = cs.create_envelope(envelope_name)

    print("[*] 2. Upload do documento...")
    document_id = cs.upload_document_base64(envelope_id, doc_path)

    print(
        "[*] 3. Ordem Consultor → Coordenador → Cliente → Diretor — "
        f"{len(sign_sequence)} signatário(s) ativo(s)...",
    )
    for om in omitidos or []:
        print(
            f"    > {om['papel']} ({om['email']}) omitido — {om.get('motivo', 'e-mail duplicado')}.",
        )
    first_group = min((p.get("group", 1) for p in sign_sequence), default=1)
    first_group_signer_ids: list[str] = []

    for i, person in enumerate(sign_sequence):
        group_num = person.get("group", i + 1)
        signer_id = cs.create_signer(
            envelope_id, person["name"], person["email"], group_num
        )
        cs.create_sign_requirement(envelope_id, signer_id, document_id)
        cs.create_auth_requirement(envelope_id, signer_id, document_id)

        if group_num == first_group:
            first_group_signer_ids.append(signer_id)
        rotulo = person.get("papel") or person["name"]
        print(
            f"    > {rotulo} ({person['email']}) adicionado ao Grupo {group_num}."
        )

    print("[*] 4. Ativando envelope...")
    cs.activate_envelope(envelope_id)

    if first_group_signer_ids:
        print(f"[*] Disparando notificação para o Grupo {first_group}...")
        for signer_id in first_group_signer_ids:
            cs.notify_signer_manual(envelope_id, signer_id)

    print(f"[v] Envelope {envelope_id} ativado! Fluxo sequencial automático iniciado.")
    return envelope_id


def _disparar_aviso_comercial_etapa1(
    deal: dict, numeros_pedidos: dict[str, str] | None
) -> None:
    """E-mail informativo ao comercial (não assina); junto com notificação do grupo 1."""
    try:
        enviar_aviso_comercial_etapa1(deal, numeros_pedidos)
    except Exception as exc:
        deal_id = str(deal.get("id", "")).strip()
        print(
            f"[!] Deal {deal_id}: falha ao preparar aviso comercial informativo: {exc}",
            flush=True,
        )


# ---------- GERAÇÃO DO CONTRATO ----------
FONTE_CONTRATO = "Calibri"


def _definir_fonte_run_calibri(run) -> None:
    """Garante Calibri no run (docxtpl pode herdar Times New Roman do template)."""
    if not run.text:
        return
    run.font.name = FONTE_CONTRATO
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        r_fonts.set(qn(attr), FONTE_CONTRATO)


def _iter_runs_paragrafo(para):
    for child in para._element:
        if child.tag == qn("w:hyperlink"):
            for r_elem in child.findall(qn("w:r")):
                yield Run(r_elem, para)
        elif child.tag == qn("w:r"):
            yield Run(child, para)


def _iter_blocos_texto_documento(document):
    for para in document.paragraphs:
        yield para
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para
    for section in document.sections:
        for bloco in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
        ):
            if bloco is None:
                continue
            for para in bloco.paragraphs:
                yield para
            for table in bloco.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            yield para


def aplicar_fonte_calibri_documento(document) -> None:
    """Normaliza a fonte de todo texto gerado no contrato (inclui hyperlinks)."""
    for para in _iter_blocos_texto_documento(document):
        for run in _iter_runs_paragrafo(para):
            _definir_fonte_run_calibri(run)


def fill_contract(
    deal_data,
    *,
    numeros_pedidos: dict[str, str] | None = None,
    template_path: str | None = None,
):
    modelo = template_path or MODELO_DOCX
    if not os.path.exists(modelo):
        print(f"[!] Erro: Modelo '{modelo}' não encontrado.")
        return None
    try:
        doc = DocxTemplate(modelo)
    except Exception as e:
        print(f"[!] Erro ao abrir docx: {e}")
        return None

    qtd_sole = get_val(deal_data, FIELD_QTD_SOLE)
    sole_consultoria = get_val(deal_data, FIELD_QUALIDADE_ENERGIA)
    qualidade_energia_uc = get_val(deal_data, FIELD_INDICADORES_QUALIDADE)
    gestao_acl = get_val(deal_data, FIELD_GESTAO_ACL)
    gestao_usina = get_val(deal_data, FIELD_GESTAO_USINA_FOTOVOLTAICA)
    percentual_exito = get_enum_label(deal_data, FIELD_PERCENTUAL_EXITO).strip()
    if not percentual_exito:
        percentual_exito = "A definir"
    inscricao_estadual = get_val(deal_data, FIELD_INSCRICAO_ESTADUAL).strip()
    inscricao_municipal = get_val(deal_data, FIELD_INSCRICAO_MUNICIPAL).strip()

    raw_dt_implantacao = get_val(deal_data, FIELD_DATA_IMPLANTACAO)
    dt_implantacao = (
        formatar_data_ptbr(raw_dt_implantacao) if raw_dt_implantacao else "A definir"
    )

    raw_dt_primeira_cobranca = get_val(deal_data, FIELD_DATA_PRIMEIRA_COBRANCA)
    dt_primeira_cobranca = (
        formatar_data_ptbr(raw_dt_primeira_cobranca)
        if raw_dt_primeira_cobranca
        else "A definir"
    )

    documento_pipe = get_val(deal_data, FIELD_DOCUMENTO)
    contratante = get_nome_cliente(deal_data)
    parceiro_plune = None
    try:
        parceiro_plune = buscar_parceiro_plune_por_documento(documento_pipe, contratante)
    except PluneError as e:
        print(f"[!] Aviso: busca Plune por CNPJ falhou ({e}); usando dados do Pipedrive no contrato.")

    if parceiro_plune:
        nome_cliente = parceiro_plune.get("razao_social") or contratante
        nome_fantasia = (parceiro_plune.get("nome_fantasia") or "").strip()
        endereco = parceiro_plune.get("endereco") or get_val(deal_data, FIELD_ENDERECO)
        cidade = (parceiro_plune.get("cidade") or "").strip()
        if not cidade:
            cidade, _ = get_cidade_estado(deal_data)
        estado = (parceiro_plune.get("uf") or "").strip()
        if not estado:
            _, estado = get_cidade_estado(deal_data)
        documento = parceiro_plune.get("documento_formatado") or documento_pipe
        print(f"[*] Contrato com dados do Plune: {nome_cliente} ({documento})")
    else:
        # Cliente ainda não no Plune: razão social e nome fantasia = Contratante (Pipedrive)
        nome_cliente = contratante
        nome_fantasia = contratante
        endereco = get_val(deal_data, FIELD_ENDERECO)
        cidade, estado = get_cidade_estado(deal_data)
        documento = documento_pipe

    if numeros_pedidos is None:
        deal_id = str(deal_data.get("id", "")).strip()
        numeros_pedidos = {}
        if deal_id:
            try:
                numeros_pedidos = obter_numeros_pedidos_plune_deal(deal_id)
            except PluneError as exc:
                print(
                    f"[!] Aviso: não foi possível buscar pedidos Plune para o contrato "
                    f"(deal {deal_id}): {exc}",
                )

    contexto = {
        "numero_contrato": get_numero_contrato(deal_data),
        "numeros_pedidos": formatar_linha_pedidos_plune_contrato(numeros_pedidos),
        "nome_cliente": nome_cliente,
        "nome_fantasia": nome_fantasia,
        "endereco": endereco,
        "cidade": cidade,
        "estado": estado,
        "documento": documento,
        "sole_web": formatar_quantidade_uc(qtd_sole),
        "sole_consultoria": formatar_quantidade_uc(sole_consultoria),
        "gestao_acl": formatar_quantidade_uc(gestao_acl),
        "gestao_usina_fotovoltaica": formatar_quantidade_uc(gestao_usina),
        "qualidade_energia": formatar_quantidade_uc(qualidade_energia_uc),
        "percentual_exito": percentual_exito,
        "inscricao_estadual": inscricao_estadual,
        "inscricao_municipal": inscricao_municipal,
        "valor_mensal": formatar_moeda(get_val(deal_data, FIELD_VALOR_MENSAL)),
        "valor_implantacao": formatar_moeda(
            get_val(deal_data, FIELD_VALOR_IMPLANTACAO)
        ),
        "data_hoje": datetime.now().strftime("%d/%m/%Y"),
        "data_inicio": formatar_data_ptbr(
            deal_data.get("won_time") or deal_data.get("add_time")
        ),
        "data_pagamento_implantacao": dt_implantacao,
        "data_pagamento_primeira_cobranca": dt_primeira_cobranca,
        "contato_gestor": get_contato_gestor_contrato(deal_data),
        "contato_financeiro": get_val(deal_data, FIELD_CONTATO_FINANCEIRO),
        "contato_contratante": get_val(deal_data, FIELD_CONTATO_CONTRATANTE),
        "notas": get_val(deal_data, FIELD_NOTAS).strip(),
    }

    clean_title = str(deal_data.get("title", "SemTitulo")).replace("/", "-")
    nome_saida = f"{PASTA_SAIDA}/Contrato_{deal_data['id']}_{clean_title}.docx"
    doc.render(contexto)
    aplicar_fonte_calibri_documento(doc.docx)
    doc.save(nome_saida)
    print(f"[v] Contrato Gerado: {nome_saida}")
    return nome_saida


_PASTA_TEMPLATES_RUNTIME = os.path.join("runtime", "templates")
_GLOB_TEMPLATE_DEAL = "template_*_{deal_id}_*.docx"


def _remover_arquivo_template_local(caminho: str) -> bool:
    path = str(caminho or "").strip()
    if not path or not os.path.isfile(path):
        return False
    try:
        os.remove(path)
        return True
    except OSError as exc:
        print(f"[!] Falha ao remover template local {path!r}: {exc}", flush=True)
        return False


def limpar_templates_locais_deal(deal_id: str, *, manter: str | None = None) -> int:
    """Remove templates baixados do Pipedrive para o deal, exceto ``manter``."""
    deal_id = str(deal_id).strip()
    if not deal_id or not os.path.isdir(_PASTA_TEMPLATES_RUNTIME):
        return 0

    manter_norm = os.path.normpath(manter) if manter else None
    pattern = os.path.join(
        _PASTA_TEMPLATES_RUNTIME, _GLOB_TEMPLATE_DEAL.format(deal_id=deal_id)
    )
    removidos = 0
    for caminho in glob.glob(pattern):
        caminho_norm = os.path.normpath(caminho)
        if manter_norm and caminho_norm == manter_norm:
            continue
        if _remover_arquivo_template_local(caminho_norm):
            removidos += 1
    return removidos


def limpar_templates_orfaos_runtime() -> int:
    """
    Remove arquivos em runtime/templates que não estão referenciados em
    envelopes_pending.template_local_path (restos de regenerações anteriores).
    """
    if not os.path.isdir(_PASTA_TEMPLATES_RUNTIME):
        return 0

    from core.database import db_conn

    ativos: set[str] = set()
    with db_conn() as conn:
        rows = conn.execute(
            """
            SELECT template_local_path
            FROM envelopes_pending
            WHERE template_local_path IS NOT NULL AND template_local_path <> ''
            """
        ).fetchall()
    for row in rows:
        path = str(row.get("template_local_path") or "").strip()
        if path:
            ativos.add(os.path.normpath(path))

    removidos = 0
    for caminho in glob.glob(os.path.join(_PASTA_TEMPLATES_RUNTIME, "template_*.docx")):
        caminho_norm = os.path.normpath(caminho)
        if caminho_norm in ativos:
            continue
        if _remover_arquivo_template_local(caminho_norm):
            removidos += 1

    if removidos:
        print(
            f"[*] Templates: {removidos} arquivo(s) órfão(s) removido(s) "
            f"de {_PASTA_TEMPLATES_RUNTIME}.",
            flush=True,
        )
    return removidos


def _salvar_template_deal_local(
    *, deal_id: str, won_time: str, file_id: int, nome: str, conteudo: bytes
) -> str:
    os.makedirs(_PASTA_TEMPLATES_RUNTIME, exist_ok=True)
    sufixo = hashlib.sha1(str(won_time).encode("utf-8")).hexdigest()[:10]
    base = os.path.splitext(os.path.basename(nome or "contrato_padrao.docx"))[0]
    safe_base = base.replace("/", "-").replace("\\", "-").strip() or "contrato_padrao"
    caminho = os.path.join(
        _PASTA_TEMPLATES_RUNTIME,
        f"template_{safe_base}_{deal_id}_{file_id}_{sufixo}.docx",
    )
    with open(caminho, "wb") as f:
        f.write(conteudo)
    return caminho


def _fluxo_hub_pre_aprovacao(
    deal_id: str,
    parceiro_criado_plune: bool,
    *,
    numeros_pedidos: dict[str, str] | None = None,
) -> None:
    """
    Antes da aprovação Plune (ganho / regeneração com troca de contrato e valores).

    - Dev (DEV_HUB_SEM_APROVACAO_PLUNE=true): cria na 1ª vez e atualiza a cada regeneração.
    - Produção: não cria; só atualiza se o pedido HUB já existir (ex.: teste com flag dev).
      Criação em processar_contratos_assinados, após aprovar_pedidos_plune.
    """
    if numeros_pedidos is not None and not any(
        str(v).strip() for v in numeros_pedidos.values() if v is not None
    ):
        print(
            f"[*] HUB: sem números de pedido Plune — deal {deal_id} (aguardar Plune).",
            flush=True,
        )
        return
    garantir_envelope_para_hub(
        deal_id, parceiro_plune_criado=parceiro_criado_plune
    )
    from core.hub_pedido import tentar_sincronizar_pedido_hub_deal

    tentar_sincronizar_pedido_hub_deal(
        deal_id,
        parceiro_plune_criado=parceiro_criado_plune,
        permitir_criacao=DEV_HUB_SEM_APROVACAO_PLUNE and not PULAR_HUB,
    )


def log_resultado_plune(deal_id: str, result: dict) -> None:
    pedidos = result.get("pedidos")
    if not pedidos:
        status = result.get("status")
        if status == "created":
            print(
                f"[v] Plune: pedido criado — id={result.get('pedido_id')} "
                f"(deal {deal_id} registrado no MySQL).",
                flush=True,
            )
        elif status == "skipped":
            print(
                f"[*] Plune: ignorado para deal {deal_id} — {result} "
                f"(já consta no MySQL ou já existe pedido com PedidoIntegracao).",
                flush=True,
            )
        else:
            print(f"[v] Plune: resultado {result}", flush=True)
        return

    print(
        f"[v] Plune: resultado geral para deal {deal_id}: {result.get('status')}",
        flush=True,
    )
    for pedido in pedidos:
        tipo = pedido.get("tipo")
        status = pedido.get("status")
        pedido_id = pedido.get("pedido_id")
        integracao = pedido.get("pedido_integracao")
        reason = pedido.get("reason")
        aprovado = pedido.get("aprovado")
        extra = f", motivo={reason}" if reason else ""
        if aprovado is not None:
            extra += f", aprovado={aprovado}"
        print(
            f"    -> {tipo}: {status}, id={pedido_id or '-'}, integracao={integracao}{extra}",
            flush=True,
        )


def _reabrir_deal_apos_falha(deal_id: str, mensagem: str) -> None:
    print(
        f"[!] Deal {deal_id}: reabrindo card no Pipedrive (etapa Negociação).",
        flush=True,
    )
    print(f"    - {mensagem}", flush=True)
    reabrir_deal_falha_automacao(deal_id, [mensagem])
    print(
        f"[v] Deal {deal_id} reaberto no Pipedrive. Não foi gravado em deals_processed; "
        "após corrigir, mova novamente para Contrato.",
        flush=True,
    )


def criar_pedidos_plune_no_ganho(deal_id: str, contexto_envio: str) -> None:
    print(
        f"    -> Criando/garantindo dois pedidos no Plune ({contexto_envio}; deal {deal_id})...",
        flush=True,
    )
    try:
        result = criar_pedido_plune(deal_id)
        log_resultado_plune(deal_id, result)
    except DealValidationError as validation_error:
        print(
            f"[!] Deal {deal_id}: parceiro novo no Plune exige CEP no Pipedrive; reabrindo card.",
            flush=True,
        )
        for erro in validation_error.mensagens:
            print(f"    - {erro}", flush=True)
        reabrir_deal_com_erros(deal_id, validation_error.mensagens)
        print(
            f"[*] Se o deal {deal_id} já estava em deals_processados, remova o estado com "
            f"`python scripts/automacao_db.py rm deal {deal_id} -y` antes de mover para Contrato de novo.",
            flush=True,
        )
    except PluneError as exc:
        print(
            f"[!] Plune: FALHA ao criar/garantir pedidos para deal {deal_id}: {exc}",
            flush=True,
        )
        _reabrir_deal_apos_falha(deal_id, f"Plune: {exc}")


def _envelope_clicksign_real(envelope_id: str | None) -> bool:
    eid = str(envelope_id or "").strip()
    return bool(eid) and not eid.startswith("sem-envelope-")


def _retomar_fluxo_interrompido(registro: dict | None) -> bool:
    """
    Retoma Plune/contrato quando há estado em envelopes_pending mas deals_processed
    ainda não foi gravado (falha parcial ou queda do processo).

    Poll com deal já em deals_processed é ignorado antes deste ponto.
    Para reprocessar do zero: `automacao_db rm deal <id>` e reenviar o formulário.
    """
    return registro is not None


def processar_deals_pendentes():
    """Processa deals enfileirados pelo formulário web (deal_forms validated/submitted)."""
    if not AUTOMACAO_WORKER_ENABLED:
        print(
            "[*] AUTOMACAO_WORKER_ENABLED=0 — worker de formulário desligado "
            "(sem Plune/contrato/e-mail comercial).",
            flush=True,
        )
        return
    if not FORMULARIO_WEB_ENABLED:
        print(
            "[*] FORMULARIO_WEB_ENABLED=0 — worker de formulário desligado.",
            flush=True,
        )
        return
    limpar_templates_orfaos_runtime()
    ids_processados, ids_legado = carregar_deals_processados()
    deal_ids = listar_deal_ids_formulario_aguardando_worker(
        exclude_deal_ids=ids_processados | ids_legado
    )
    if not deal_ids:
        print(
            "[*] Formulário web: nenhum deal validated aguardando automação.",
            flush=True,
        )
        return

    n_sem_deal = n_fechado = n_fora_contrato = n_legado = n_dup = n_sem_evento = n_data = 0
    algum_processado = False

    for deal_id_int in deal_ids:
        deal_id = str(deal_id_int)

        if deal_id in ids_legado:
            n_legado += 1
            if deal_id not in _deals_legado_ja_avisados:
                _deals_legado_ja_avisados.add(deal_id)
                print(
                    f"[*] Deal {deal_id}: ignorado — está em deals_legacy_block (MySQL).",
                    flush=True,
                )
            continue

        if deal_id in ids_processados:
            n_dup += 1
            continue

        deal = buscar_deal_por_id(deal_id)
        if not deal:
            n_sem_deal += 1
            print(f"[!] Deal {deal_id}: não encontrado no Pipedrive.", flush=True)
            continue

        if deal.get("status") != "open":
            n_fechado += 1
            print(
                f"[*] Deal {deal_id}: ignorado — card não está aberto no Pipedrive "
                f"(status={deal.get('status')!r}).",
                flush=True,
            )
            continue

        if not deal_elegivel_formulario_contrato(deal):
            n_fora_contrato += 1
            print(
                f"[*] Deal {deal_id}: ignorado — card não está na etapa Contrato "
                f"(stage_id={deal.get('stage_id')!r}).",
                flush=True,
            )
            continue

        evento_str = deal.get("update_time") or deal.get("add_time")
        if not evento_str:
            n_sem_evento += 1
            continue

        registro = buscar_por_deal_id(deal_id)

        if _parse_deal_timestamp_utc(evento_str) is None:
            n_data += 1
            print(
                f"[!] Deal {deal_id}: update_time em formato não reconhecido: {evento_str!r}",
                flush=True,
            )
            continue

        print(
            f"\n[!] FORMULÁRIO ENVIADO — automação disparada! "
            f"Deal ID: {deal_id} | Título: {deal.get('title')}",
            flush=True,
        )

        try:
            atualizar_deal_form_status(deal_id, "processing")
        except Exception as exc:
            print(
                f"[!] Deal {deal_id}: falha ao marcar formulário como processing: {exc}",
                flush=True,
            )

        prep = preparar_deal_para_automacao(
            deal, formulario_web_enabled=FORMULARIO_WEB_ENABLED
        )
        if prep.skipped_reason:
            status_info = (
                f" (status={prep.form_status})" if prep.form_status else ""
            )
            print(
                f"[*] Deal {deal_id}: ignorado — {prep.skipped_reason}{status_info}.",
                flush=True,
            )
            continue

        deal = prep.deal
        print(
            f"[*] Deal {deal_id}: usando payload do formulário web (merge form > Pipe).",
            flush=True,
        )

        try:
            # Idempotência: deal em deals_processed é ignorado no início do loop.
            # Retomada: envelope_pending sem deals_processed (falha parcial / queda do processo).
            # Reprocessar do zero: `automacao_db rm deal <id>` e reenviar o formulário.
            if registro and registro.get("pedidos_plune_aprovados"):
                print(
                    f"[*] Deal {deal_id}: ignorado — pedidos Plune já aprovados "
                    "(regra: sem regeneração após última assinatura).",
                    flush=True,
                )
                _registrar_deal_processado_worker(deal_id, evento_str)
                ids_processados.add(deal_id)
                algum_processado = True
                continue

            # Plune: criar ou retomar (estado em envelopes_pending sem deals_processed)
            numeros_pedidos: dict[str, str] = {}
            parceiro_criado_plune = False
            if _retomar_fluxo_interrompido(registro):
                print(
                    f"    -> Retomada: atualizando pedidos no Plune (deal {deal_id})...",
                    flush=True,
                )
                try:
                    from core.plune_pedido import atualizar_pedidos_plune

                    result_plune = atualizar_pedidos_plune(deal_id)
                    log_resultado_plune(deal_id, result_plune)
                    numeros_pedidos = extrair_numeros_pedidos_plune(result_plune)
                    _fluxo_hub_pre_aprovacao(
                        deal_id,
                        bool(registro.get("parceiro_plune_criado")),
                        numeros_pedidos=numeros_pedidos,
                    )
                except DealValidationError as validation_error:
                    print(
                        f"[!] Deal {deal_id}: parceiro novo no Plune exige CEP no Pipedrive; reabrindo card.",
                        flush=True,
                    )
                    for erro in validation_error.mensagens:
                        print(f"    - {erro}", flush=True)
                    reabrir_deal_com_erros(deal_id, validation_error.mensagens)
                    algum_processado = True
                    continue
                except PluneError as exc:
                    print(
                        f"[!] Plune: FALHA ao atualizar pedidos (deal {deal_id}): {exc}",
                        flush=True,
                    )
                    _reabrir_deal_apos_falha(deal_id, f"Plune: {exc}")
                    algum_processado = True
                    continue
            else:
                print(
                    f"    -> Criando pedidos no Plune (antes do contrato)...",
                    flush=True,
                )
                try:
                    result_plune = criar_pedido_plune(deal_id, anexar_contrato=False)
                    parceiro_criado_plune = bool(result_plune.get("parceiro_criado"))
                    log_resultado_plune(deal_id, result_plune)
                    # Se o MySQL local já marca o pedido como criado, mas o deal não tem
                    # registro de envelope pendente (registro=None), tratamos como
                    # regeneração e aplicamos atualização no Plune.
                    pedidos = result_plune.get("pedidos") or []
                    if any(
                        (p.get("reason") == "already_in_local_state") for p in pedidos
                    ):
                        print(
                            f"    -> Detectado pedido já criado localmente; atualizando pedidos no Plune (deal {deal_id})...",
                            flush=True,
                        )
                        from core.plune_pedido import atualizar_pedidos_plune

                        result_plune = atualizar_pedidos_plune(deal_id)
                        log_resultado_plune(deal_id, result_plune)
                    numeros_pedidos = extrair_numeros_pedidos_plune(result_plune)
                    _fluxo_hub_pre_aprovacao(
                        deal_id,
                        parceiro_criado_plune,
                        numeros_pedidos=numeros_pedidos,
                    )
                except DealValidationError as validation_error:
                    print(
                        f"[!] Deal {deal_id}: parceiro novo no Plune exige CEP no Pipedrive; reabrindo card.",
                        flush=True,
                    )
                    for erro in validation_error.mensagens:
                        print(f"    - {erro}", flush=True)
                    reabrir_deal_com_erros(deal_id, validation_error.mensagens)
                    algum_processado = True
                    continue
                except PluneError as exc:
                    print(
                        f"[!] Plune: FALHA ao criar pedidos antes do contrato (deal {deal_id}): {exc}",
                        flush=True,
                    )
                    _reabrir_deal_apos_falha(deal_id, f"Plune: {exc}")
                    algum_processado = True
                    continue

            print(f"    -> Gerando contrato...", flush=True)

            template_path = None
            template_file_id = None
            try:
                tpl = baixar_docx_contrato_padrao_deal(deal_id)
                if tpl:
                    conteudo_tpl, nome_tpl, file_id = tpl
                    template_file_id = file_id
                    template_path = _salvar_template_deal_local(
                        deal_id=deal_id,
                        won_time=evento_str,
                        file_id=file_id,
                        nome=nome_tpl,
                        conteudo=conteudo_tpl,
                    )
                    removidos = limpar_templates_locais_deal(
                        deal_id, manter=template_path
                    )
                    if removidos:
                        print(
                            f"[*] Deal {deal_id}: {removidos} template(s) antigo(s) "
                            f"removido(s) de runtime/templates.",
                            flush=True,
                        )
                    print(
                        f"[*] Deal {deal_id}: usando template do Pipedrive "
                        f"(file_id={file_id}, nome={nome_tpl!r}).",
                        flush=True,
                    )
            except Exception as exc:
                print(
                    f"[!] Deal {deal_id}: falha ao baixar template contrato_padrao do Pipedrive "
                    f"(seguindo com template local): {exc}",
                    flush=True,
                )

            doc_path = fill_contract(
                deal,
                numeros_pedidos=numeros_pedidos,
                template_path=template_path,
            )

            sequence_dinamica = extrair_signatarios(deal)

            if not doc_path:
                print(
                    f"[!] Deal {deal_id}: falha ao gerar contrato (.docx).",
                    flush=True,
                )
                _reabrir_deal_apos_falha(
                    deal_id, "Falha ao gerar o arquivo do contrato (.docx)."
                )
                algum_processado = True
                continue

            if doc_path:

                def _apos_contrato_dev_sem_envelope_msg():
                    if not TESTE_PLUNE_SEM_ASSINATURA:
                        print(
                            f"[*] Pós-assinatura: com DEV_PULAR_CLICKSIGN não há envelope para aprovar depois. "
                            f"Use TESTE_PLUNE_SEM_ASSINATURA=true em dev ou desligue o pulo para criar envelope.",
                            flush=True,
                        )

                if DEV_PULAR_CLICKSIGN:
                    print(
                        f"    -> [DEV] DEV_PULAR_CLICKSIGN=1 — contrato gravado em «{doc_path}»; "
                        f"API Clicksign não será chamada.",
                        flush=True,
                    )
                    parceiro_para_envelope = parceiro_criado_plune
                    if registro:
                        parceiro_para_envelope = bool(
                            registro.get("parceiro_plune_criado")
                        ) or parceiro_criado_plune
                    salvar_envelope_pendente(
                        deal_id,
                        f"sem-envelope-{deal_id}",
                        f"Deal {deal_id} (DEV sem Clicksign)",
                        template_file_id=template_file_id,
                        template_local_path=template_path,
                        parceiro_plune_criado=parceiro_para_envelope,
                    )
                    _registrar_deal_processado_worker(deal_id, evento_str)
                    ids_processados.add(deal_id)
                    algum_processado = True
                    print(
                        f"[v] Deal {deal_id} registrado no MySQL (DEV sem Clicksign).",
                        flush=True,
                    )
                    try:
                        _disparar_aviso_comercial_etapa1(deal, numeros_pedidos)
                    except Exception as exc:
                        print(
                            f"[!] Deal {deal_id}: aviso comercial falhou (fluxo já registrado): {exc}",
                            flush=True,
                        )
                    _apos_contrato_dev_sem_envelope_msg()
                elif len(sequence_dinamica) > 0:
                    print(f"    -> Enviando para Clicksign...", flush=True)
                    envelope_name = f"Contrato Deal {deal_id} - {deal.get('title')}"

                    # Regeneração: cancela envelope Clicksign real antes de criar o novo.
                    envelope_antigo = (
                        str(registro["envelope_id"])
                        if registro and registro.get("envelope_id")
                        else ""
                    )
                    if _envelope_clicksign_real(envelope_antigo):
                        try:
                            cs_tmp = ClicksignClient(
                                CLICKSIGN_BASE_URL, CLICKSIGN_ACCESS_TOKEN
                            )
                            cs_tmp.cancel_envelope(envelope_antigo)
                        except Exception as exc:
                            print(
                                f"[!] Clicksign: falha ao tentar cancelar envelope antigo "
                                f"{registro.get('envelope_id')}: {exc}",
                                flush=True,
                            )

                    envelope_id = clicksign_fire_and_forget(
                        doc_path,
                        envelope_name,
                        sequence_dinamica,
                        omitidos=signatarios_omitidos_por_email_duplicado(
                            deal, sequence_dinamica
                        ),
                    )

                    parceiro_para_envelope = parceiro_criado_plune
                    if registro:
                        parceiro_para_envelope = bool(
                            registro.get("parceiro_plune_criado")
                        ) or parceiro_criado_plune
                    salvar_envelope_pendente(
                        deal_id,
                        envelope_id,
                        envelope_name,
                        template_file_id=template_file_id,
                        template_local_path=template_path,
                        parceiro_plune_criado=parceiro_para_envelope,
                    )
                    _registrar_deal_processado_worker(deal_id, evento_str)
                    ids_processados.add(deal_id)
                    algum_processado = True
                    print(
                        f"[v] Contrato enviado ao Clicksign. Deal {deal_id} registrado no MySQL.",
                        flush=True,
                    )
                    try:
                        _disparar_aviso_comercial_etapa1(deal, numeros_pedidos)
                    except Exception as exc:
                        print(
                            f"[!] Deal {deal_id}: aviso comercial falhou (fluxo já registrado): {exc}",
                            flush=True,
                        )

                    if not TESTE_PLUNE_SEM_ASSINATURA:
                        print(
                            f"[*] Pós-assinatura: quando o envelope fechar, os pedidos Plune serão aprovados. "
                            f"Pendências de envelope no MySQL (envelopes_pending).",
                            flush=True,
                        )
                else:
                    print(
                        "[!] Nenhum e-mail de signatário encontrado no Deal. Contrato não enviado para assinatura.",
                        flush=True,
                    )
                    _reabrir_deal_apos_falha(
                        deal_id,
                        "Nenhum e-mail de signatário encontrado no deal; "
                        "contrato não enviado para assinatura.",
                    )
                    algum_processado = True

        except Exception as e:
            print(f"[!] Erro ao processar Deal {deal_id}: {e}", flush=True)
            traceback.print_exc()
            try:
                _reabrir_deal_apos_falha(deal_id, str(e))
                algum_processado = True
            except Exception as reopen_exc:
                print(
                    f"[!] Deal {deal_id}: falha ao reabrir card após erro: {reopen_exc}",
                    flush=True,
                )

    if not algum_processado:
        partes = []
        if n_sem_deal:
            partes.append(f"{n_sem_deal} deal(s) não encontrados no Pipe")
        if n_fechado:
            partes.append(f"{n_fechado} card(s) fechados no Pipe")
        if n_fora_contrato:
            partes.append(f"{n_fora_contrato} fora da etapa Contrato")
        if n_sem_evento:
            partes.append(f"{n_sem_evento} sem update_time")
        if n_legado:
            partes.append(f"{n_legado} bloqueado(s) em deals_legacy_block")
        if n_dup:
            partes.append(f"{n_dup} deal(s) já processados (deals_processed)")
        if n_data:
            partes.append(f"{n_data} update_time inválido")
        if partes:
            print(
                f"[*] Ciclo formulário: {', '.join(partes)} — nenhum fluxo novo disparado.",
                flush=True,
            )


def processar_contratos_assinados():
    """Quando o envelope Clicksign fechou (todos assinaram), aprova os pedidos no Plune."""
    if TESTE_PLUNE_SEM_ASSINATURA:
        return

    pendentes = listar_aguardando_pedido_plune()
    if not pendentes:
        print(
            f"[*] Plune (pós-assinatura): nenhum envelope aguardando pedido; "
            f"pedidos criados localmente: {len(carregar_pedidos_plune_criados())} (MySQL).",
            flush=True,
        )
        return

    print(
        f"[*] Plune (pós-assinatura): {len(pendentes)} envelope(s) aguardando aprovação; "
        f"pedidos criados localmente: {len(carregar_pedidos_plune_criados())}.",
        flush=True,
    )

    cs = ClicksignClient(CLICKSIGN_BASE_URL, CLICKSIGN_ACCESS_TOKEN)

    for record in pendentes:
        deal_id = str(record.get("deal_id"))
        envelope_id = record.get("envelope_id")
        status = cs.get_envelope_status(envelope_id)
        if status != "closed":
            print(
                f"[*] Deal {deal_id}: envelope {envelope_id} status={status!r} (aguardando assinaturas).",
                flush=True,
            )
            continue

        data_entrega = cs.get_ultima_assinatura_data_ptbr(envelope_id)
        print(
            f"\n[!] Contrato assinado! Deal {deal_id} → aprovando pedidos no Plune "
            f"(Data de entrega: {data_entrega})…",
            flush=True,
        )
        try:
            result = aprovar_pedidos_plune(deal_id, data_entrega=data_entrega)
            log_resultado_plune(deal_id, result)
            if result.get("status") in ("approved", "skipped"):
                marcar_pedidos_aprovados(deal_id)
                try:
                    marcar_deal_como_ganho(deal_id)
                    print(
                        f"[*] Pipedrive: deal {deal_id} marcado como ganho "
                        f"(após assinatura completa).",
                        flush=True,
                    )
                except Exception as exc:
                    print(
                        f"[!] Pipedrive: falha ao marcar deal {deal_id} como ganho: {exc}",
                        flush=True,
                    )
                from core.hub_pedido import tentar_criar_pedido_hub_deal

                registro_hub = buscar_por_deal_id(deal_id)
                tentar_criar_pedido_hub_deal(
                    deal_id,
                    parceiro_plune_criado=(
                        bool(registro_hub.get("parceiro_plune_criado"))
                        if registro_hub
                        else None
                    ),
                )
                # Cleanup do template local usado para gerar o contrato (somente após anexar assinado).
                try:
                    path_tpl = str(record.get("template_local_path") or "").strip()
                    _remover_arquivo_template_local(path_tpl)
                    limpar_templates_locais_deal(deal_id)
                    if path_tpl:
                        limpar_template_local_envelope(deal_id)
                except Exception as exc:
                    print(
                        f"[!] Deal {deal_id}: falha ao remover template local {record.get('template_local_path')!r}: {exc}",
                        flush=True,
                    )
        except PluneError as exc:
            print(
                f"[!] Plune: falha ao aprovar pedidos (deal {deal_id}): {exc}",
                flush=True,
            )
            try:
                _reabrir_deal_apos_falha(deal_id, f"Plune (pós-assinatura): {exc}")
            except Exception as reopen_exc:
                print(
                    f"[!] Deal {deal_id}: falha ao reabrir card após erro Plune: {reopen_exc}",
                    flush=True,
                )


def main():
    if not os.path.exists(PASTA_SAIDA):
        os.makedirs(PASTA_SAIDA)

    agora_utc = datetime.now(timezone.utc)

    print("--- INICIANDO AUTOMACAO (FORMULARIO WEB -> CLICKSIGN -> PLUNE) ---")
    print(
        f"[*] Script iniciado em: {formatar_data_hora_brasilia(agora_utc)} "
        f"(horário de Brasília)"
    )
    from core.config import MYSQL_DATABASE, MYSQL_HOST

    print(
        f"[*] Estado persistido em MySQL: {MYSQL_HOST}/{MYSQL_DATABASE} "
        f"(deal_forms, deals, pedidos Plune, envelopes)."
    )
    print(
        "[*] Gatilho: formulário web enviado (deal_forms validated/submitted) — "
        "não dispara ao mover o card para Contrato."
    )
    if not AUTOMACAO_WORKER_ENABLED:
        print(
            "[*] AUTOMACAO_WORKER_ENABLED=0 — fila de formulário desligada "
            "(sem Plune/contrato/e-mail comercial)."
        )
    elif not FORMULARIO_WEB_ENABLED:
        print("[*] FORMULARIO_WEB_ENABLED=0 — fila de formulário desligada.")
    print(
        f"[*] Aviso comercial (etapa 1) → {EMAIL_COMERCIAL_AUTOMACAO} "
        f"(core/gebras_defaults.py)."
    )
    if TESTE_PLUNE_SEM_ASSINATURA:
        print(
            "[*] MODO TESTE: pedido Plune logo após Clicksign (processar_contratos_assinados fica em silêncio de propósito)."
        )
    if DEV_PULAR_CLICKSIGN:
        print(
            "[*] DEV_PULAR_CLICKSIGN=1 — a API Clicksign não é chamada (só gera o .docx no disco; use com TESTE_PLUNE_SEM_ASSINATURA em dev para ir ao Plune)."
        )
    if PULAR_HUB:
        print(
            "[*] PULAR_HUB=1 — criação de pedido no HUB desligada "
            "(atualizações e validações HUB continuam)."
        )
    elif DEV_HUB_SEM_APROVACAO_PLUNE:
        print(
            "[*] DEV_HUB_SEM_APROVACAO_PLUNE=1 — pedido HUB após Plune no envio do formulário "
            "(sem esperar aprovação Plune pós-assinatura)."
        )
    print(
        f"[*] Poll a cada {INTERVALO_POLLING_SEGUNDOS}s — cada ciclo imprime um resumo se nada novo ocorrer.\n"
    )

    try:
        n_ciclo = 0
        while True:
            n_ciclo += 1
            agora = formatar_data_hora_brasilia(datetime.now(timezone.utc))
            print(f"--- Ciclo #{n_ciclo} ({agora} — Brasília) ---", flush=True)
            processar_deals_pendentes()
            processar_contratos_assinados()
            time.sleep(INTERVALO_POLLING_SEGUNDOS)
    except KeyboardInterrupt:
        print("\n[!] Automação encerrada pelo usuário.")


if __name__ == "__main__":
    main()
